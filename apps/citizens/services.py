import openpyxl
from openpyxl.styles import Font, Alignment
from openpyxl.comments import Comment
from openpyxl.utils import get_column_letter

import io
import datetime
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor
from django.contrib.auth import get_user_model
from django.db import transaction
from .models import CitizenProfile, Skill, Category, Experience, Education

User = get_user_model()

def generate_excel_template():
    wb = openpyxl.Workbook()
    
    # --- 1. INSTRUCTIONS ---
    ws_instructions = wb.active
    ws_instructions.title = "INSTRUCTIONS"
    
    instructions_text = [
        ["GUIDE D'IMPORTATION DES UTILISATEURS ET PROFILS"],
        [""],
        ["Règles importantes :"],
        ["1. L'identifiant 'Username' est la CLÉ qui relie toutes les feuilles entre elles."],
        ["2. Cet identifiant DOIT être unique et identique pour un même utilisateur dans toutes les feuilles."],
        ["3. Ne modifiez pas l'ordre des colonnes ni leurs en-têtes (la première ligne de chaque feuille)."],
        ["4. Les dates doivent être au format AAAA-MM-JJ (ex: 2026-05-15)."],
        ["5. Les champs 'True/False' (Vrai/Faux) doivent respecter cette syntaxe exacte."],
        ["6. Supprimez les lignes d'exemple avant d'importer vos données."],
        [""],
        ["Comment remplir chaque feuille ?"],
        ["- USERS : Crée l'utilisateur. Role possibles: 'soutien', 'expert_ca', 'expert_hca', 'talent', 'diaspora', 'admin'."],
        ["- PROFILES : Ajoute les détails du profil citoyen. Le 'Location' doit être l'un des départements du Bénin."],
        ["- SKILLS : Vous pouvez ajouter plusieurs compétences par utilisateur (1 ligne = 1 compétence)."],
        ["- EXPERIENCES : 1 ligne = 1 expérience. Si 'En cours' est True, 'Date fin' peut être vide."],
        ["- EDUCATION : 1 ligne = 1 diplôme ou formation. Date fin peut être vide si en cours."]
    ]
    
    for row in instructions_text:
        ws_instructions.append(row)
        
    title_font = Font(bold=True, size=14, color="003366")
    ws_instructions['A1'].font = title_font
    
    # --- 2. USERS ---
    ws_users = wb.create_sheet(title="USERS")
    headers_users = ["Username", "Email", "Password", "First Name", "Last Name", "Is Active", "Is Staff", "Role", "Phone Number", "Is Verified"]
    ws_users.append(headers_users)
    ws_users.append(["user1", "user1@email.com", "ChangeMe123!", "Jean", "Dupont", "True", "False", "talent", "+22900000000", "True"])
    
    ws_users['A1'].comment = Comment("Clé unique, sans espace", "Admin")
    ws_users['F1'].comment = Comment("True ou False", "Admin")
    ws_users['H1'].comment = Comment("'soutien', 'expert_ca', 'expert_hca', 'talent', 'diaspora', ou 'admin'", "Admin")
    
    # --- 3. PROFILES ---
    ws_profiles = wb.create_sheet(title="PROFILES")
    headers_profiles = ["Username", "Bio", "Localisation", "Disponibilité", "Titre actuel", "Années d'expérience", "Public", "Validé"]
    ws_profiles.append(headers_profiles)
    ws_profiles.append(["user1", "Développeur passionné par les solutions open-source.", "littoral", "Temps plein", "Ingénieur Logiciel", 5, "True", "False"])
    
    ws_profiles['A1'].comment = Comment("Doit correspondre au Username de la feuille USERS", "Admin")
    ws_profiles['C1'].comment = Comment("alibori, atacora, atlantique, borgou, collines, couffo, donga, littoral, mono, oueme, plateau, zou", "Admin")
    ws_profiles['G1'].comment = Comment("True ou False", "Admin")
    
    # --- 4. SKILLS ---
    ws_skills = wb.create_sheet(title="SKILLS")
    headers_skills = ["Username", "Catégorie", "Compétence"]
    ws_skills.append(headers_skills)
    ws_skills.append(["user1", "Informatique", "Python"])
    ws_skills.append(["user1", "Informatique", "Django"])
    
    # --- 5. EXPERIENCES ---
    ws_experiences = wb.create_sheet(title="EXPERIENCES")
    headers_experiences = ["Username", "Entreprise", "Poste", "Date début", "Date fin", "En cours", "Description"]
    ws_experiences.append(headers_experiences)
    ws_experiences.append(["user1", "Tech Africa", "Lead Developer", "2020-01-15", "", "True", "Développement d'applications web en Django pour le secteur public."])
    
    ws_experiences['D1'].comment = Comment("Format: YYYY-MM-DD", "Admin")
    ws_experiences['E1'].comment = Comment("Peut être vide si 'En cours' est True", "Admin")
    
    # --- 6. EDUCATION ---
    ws_education = wb.create_sheet(title="EDUCATION")
    headers_education = ["Username", "Institution", "Diplôme", "Domaine", "Date début", "Date fin"]
    ws_education.append(headers_education)
    ws_education.append(["user1", "Université d'Abomey-Calavi", "Master en Informatique", "Génie Logiciel", "2018-09-01", "2020-07-30"])
    
    ws_education['E1'].comment = Comment("Format: YYYY-MM-DD", "Admin")
    
    # --- FORMATTING (All sheets) ---
    for _ws in [ws_users, ws_profiles, ws_skills, ws_experiences, ws_education]:
        for cell in _ws[1]:
            cell.font = Font(bold=True)
            cell.alignment = Alignment(horizontal="center", vertical="center")
            
        for col in _ws.columns:
            max_length = 0
            column = col[0].column_letter
            for cell in col:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            _ws.column_dimensions[column].width = max_length + 2

    return wb


def _add_log(doc, text, status="INFO"):
    p = doc.add_paragraph()
    run = p.add_run(f"[{status}] ")
    if status == "SUCCESS":
        run.font.color.rgb = RGBColor(0, 128, 0)
    elif status == "ERROR":
        run.font.color.rgb = RGBColor(255, 0, 0)
    elif status == "WARNING":
        run.font.color.rgb = RGBColor(255, 165, 0)
    p.add_run(text)


def process_excel_import(excel_file):
    """
    Lit le fichier Excel, importe les données (Users, Profiles, Skills, etc.) 
    et retourne un rapport Word sous forme de BytesIO.
    """
    doc = DocxDocument()
    doc.add_heading(f"Rapport d'importation - {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", level=1)
    
    try:
        wb = openpyxl.load_workbook(excel_file, data_only=True)
    except Exception as e:
        _add_log(doc, f"Erreur lors de la lecture du fichier Excel : {str(e)}", "ERROR")
        return generate_word_bytes(doc)

    # Dictionnaire temporaire pour la transaction globale
    users_created = []

    try:
        with transaction.atomic():
            # 1. USERS
            if "USERS" in wb.sheetnames:
                ws_users = wb["USERS"]
                _add_log(doc, "--- Traitement de la feuille USERS ---")
                
                # S'assurer de sauter l'en-tête (min_row=2)
                for row in ws_users.iter_rows(min_row=2, values_only=True):
                    if not row[0]: # Username manquant
                        continue
                        
                    username, email, password, first_name, last_name, is_active, is_staff, role, phone_number, is_verified = row[:10]
                    
                    if User.objects.filter(username=username).exists():
                        _add_log(doc, f"L'utilisateur '{username}' existe déjà. Ignoré (mise à jour non gérée dans cet import partiel).", "WARNING")
                        continue
                        
                    user = User(
                        username=username,
                        email=email if email else '',
                        first_name=first_name if first_name else '',
                        last_name=last_name if last_name else '',
                        is_active=str(is_active).strip().lower() == 'true',
                        is_staff=str(is_staff).strip().lower() == 'true',
                        role=role if role else 'talent',
                        phone_number=phone_number if phone_number else '',
                        is_verified=str(is_verified).strip().lower() == 'true'
                    )
                    user.set_password(password if password else "DefaultPass123!")
                    user.save()
                    users_created.append(username)
                    _add_log(doc, f"Utilisateur créé: {username}", "SUCCESS")
            else:
                _add_log(doc, "Feuille 'USERS' introuvable. Arrêt de l'import.", "ERROR")
                raise ValueError("Missing USERS sheet")

            # 2. PROFILES
            if "PROFILES" in wb.sheetnames:
                ws_profiles = wb["PROFILES"]
                _add_log(doc, "--- Traitement de la feuille PROFILES ---")
                
                for row in ws_profiles.iter_rows(min_row=2, values_only=True):
                    if not row[0]: continue
                    username, bio, location, availability, current_title, years_xp, is_public, is_validated = row[:8]
                    
                    user = User.objects.filter(username=username).first()
                    if not user:
                        _add_log(doc, f"Utilisateur '{username}' introuvable pour le profil.", "ERROR")
                        continue
                        
                    profile, created = CitizenProfile.objects.get_or_create(user=user)
                    profile.bio = bio if bio else profile.bio
                    profile.location = location if location else profile.location
                    profile.availability = availability if availability else profile.availability
                    profile.current_title = current_title if current_title else profile.current_title
                    
                    try:
                        profile.years_of_experience = int(years_xp) if years_xp else 0
                    except:
                        profile.years_of_experience = 0
                        
                    profile.is_public = str(is_public).strip().lower() == 'true'
                    profile.is_validated = str(is_validated).strip().lower() == 'true'
                    profile.save()
                    _add_log(doc, f"Profil mis à jour pour: {username}", "SUCCESS")

            # 3. SKILLS
            if "SKILLS" in wb.sheetnames:
                ws_skills = wb["SKILLS"]
                _add_log(doc, "--- Traitement de la feuille SKILLS ---")
                
                for row in ws_skills.iter_rows(min_row=2, values_only=True):
                    if not row[0] or not row[2]: continue
                    username, category_name, skill_name = row[:3]
                    
                    user = User.objects.filter(username=username).first()
                    if not user: continue
                    try:
                        profile = user.citizen_profile
                    except:
                        continue
                        
                    category = None
                    if category_name:
                        category, _ = Category.objects.get_or_create(name=category_name)
                        
                    skill, _ = Skill.objects.get_or_create(name=skill_name, category=category)
                    profile.skills.add(skill)
                    _add_log(doc, f"Compétence '{skill_name}' ajoutée à '{username}'", "SUCCESS")

            # 4. EXPERIENCES
            if "EXPERIENCES" in wb.sheetnames:
                ws_experiences = wb["EXPERIENCES"]
                _add_log(doc, "--- Traitement de la feuille EXPERIENCES ---")
                
                for row in ws_experiences.iter_rows(min_row=2, values_only=True):
                    if not row[0]: continue
                    username, company, title, start_date, end_date, is_current, description = row[:7]
                    
                    user = User.objects.filter(username=username).first()
                    if not user: continue
                    try: profile = user.citizen_profile
                    except: continue
                        
                    try:
                        Experience.objects.create(
                            profile=profile,
                            company=company if company else "N/A",
                            position=title if title else "N/A",
                            start_date=start_date if start_date else datetime.date.today(),
                            end_date=end_date if (end_date and str(end_date).strip() != "") else None,
                            is_current=str(is_current).strip().lower() == 'true',
                            description=description if description else ""
                        )
                        _add_log(doc, f"Expérience '{title}' ajoutée à '{username}'", "SUCCESS")
                    except Exception as e:
                        _add_log(doc, f"Erreur date/format pour expérience de '{username}': {str(e)}", "ERROR")

            # 5. EDUCATION
            if "EDUCATION" in wb.sheetnames:
                ws_education = wb["EDUCATION"]
                _add_log(doc, "--- Traitement de la feuille EDUCATION ---")
                
                for row in ws_education.iter_rows(min_row=2, values_only=True):
                    if not row[0]: continue
                    username, institution, degree, field_of_study, start_date, end_date = row[:6]
                    
                    user = User.objects.filter(username=username).first()
                    if not user: continue
                    try: profile = user.citizen_profile
                    except: continue
                        
                    try:
                        Education.objects.create(
                            profile=profile,
                            institution=institution if institution else "N/A",
                            degree=degree if degree else "N/A",
                            field_of_study=field_of_study if field_of_study else "",
                            start_date=start_date if start_date else datetime.date.today(),
                            end_date=end_date if (end_date and str(end_date).strip() != "") else None
                        )
                        _add_log(doc, f"Formation '{degree}' ajoutée à '{username}'", "SUCCESS")
                    except Exception as e:
                        _add_log(doc, f"Erreur date/format pour formation de '{username}': {str(e)}", "ERROR")

        _add_log(doc, ">>> IMPORTATION TERMINEE AVEC SUCCES <<<", "SUCCESS")

    except Exception as e:
        _add_log(doc, f"Une erreur fatale est survenue, la transaction a été annulée: {str(e)}", "ERROR")
        
    return generate_word_bytes(doc)

def generate_word_bytes(doc):
    f = io.BytesIO()
    doc.save(f)
    f.seek(0)
    return f

def process_bulk_validation_with_report(pending_profiles, request=None):
    """
    Validates profiles, sends emails, and generates a Word report of the results.
    """
    from apps.core.utils import send_pcc_email
    from apps.core.models import Notification
    import time
    
    doc = DocxDocument()
    doc.add_heading(f"Rapport de Validation Groupée - {datetime.datetime.now().strftime('%d/%m/%Y %H:%M')}", level=1)
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Candidat'
    hdr_cells[1].text = 'Email'
    hdr_cells[2].text = 'Statut Action'
    
    _add_log(doc, f"Nombre de profils à traiter : {len(pending_profiles)}")
    _add_log(doc, "--------------------------------------------------------")
    
    success_count = 0
    error_count = 0

    for profile in pending_profiles:
        email = profile.user.email
        fullname = profile.user.get_full_name() or profile.user.username
        
        row_cells = table.add_row().cells
        row_cells[0].text = fullname
        row_cells[1].text = email
        
        try:
            with transaction.atomic():
                # 1. Validation DB
                profile.is_validated = True
                profile.save()
                
                # 2. Notification interne
                Notification.objects.create(
                    user=profile.user,
                    title="✨ Profil Validé",
                    message="Votre profil expert a été validé par l'administration. Vous pouvez maintenant postuler.",
                    link="/accounts/dashboard/"
                )
            
            # 3. Envoi de l'Email (avec délai anti-spam)
            subject = "Profil Validé - Plateforme Citoyenne des Compétences"
            context = {'user': profile.user, 'profile': profile}
            email_sent = send_pcc_email(subject, 'emails/profile_validated.html', context, [email], request=request)
            
            if email_sent:
                row_cells[2].text = "VALIDE & EMAIL ENVOYE"
                success_count += 1
            else:
                row_cells[2].text = "VALIDE (ERREUR ENVOI EMAIL)"
                error_count += 1
                
            # Pause anti-spam si ce n'est pas le dernier
            time.sleep(1)
            
        except Exception as e:
            row_cells[2].text = f"ERREUR: {str(e)}"
            error_count += 1

    doc.add_paragraph(f"\nRésumé : {success_count} réussites, {error_count} erreurs/alertes.")
    return generate_word_bytes(doc)

def process_broadcast_email_report(subject, message, target_role='all', request=None):
    """
    Sends an email to selected users with a 5-second delay and returns the Word report.
    """
    from apps.accounts.models import CustomUser
    from apps.core.utils import send_pcc_email
    import time
    
    users = CustomUser.objects.filter(is_active=True).exclude(email='')
    if target_role != 'all':
        users = users.filter(role=target_role)
    
    doc = DocxDocument()
    doc.add_heading(f"Rapport de Diffusion Générale - {datetime.datetime.now().strftime('%d/%m/%Y %H:%M')}", level=1)
    doc.add_paragraph(f"Objet : {subject}")
    doc.add_paragraph(f"Cible : {target_role}")
    doc.add_paragraph(f"Nombre de destinataires : {len(users)}")
    doc.add_paragraph("--------------------------------------------------------")
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Utilisateur'
    hdr_cells[1].text = 'Email'
    hdr_cells[2].text = 'Statut Envoi'
    
    success_count = 0
    error_count = 0

    for user in users:
        fullname = user.get_full_name() or user.username
        email = user.email
        
        row_cells = table.add_row().cells
        row_cells[0].text = fullname
        row_cells[1].text = email
        
        try:
            email_sent = send_pcc_email(
                subject=subject,
                template_name='emails/broadcast_message.html',
                context={'user': user, 'message_content': message, 'subject': subject},
                recipient_list=[email],
                request=request
            )
            
            if email_sent:
                row_cells[2].text = "ENVOYE"
                success_count += 1
            else:
                row_cells[2].text = "ECHEC (SMTP)"
                error_count += 1
                
            # DELAI DE 5 SECONDES DEMANDE PAR L'UTILISATEUR
            if len(users) > 1:
                time.sleep(5)
            
        except Exception as e:
            row_cells[2].text = f"ERREUR: {str(e)}"
            error_count += 1

    doc.add_paragraph(f"\nRésumé Final : {success_count} envoyés, {error_count} erreurs.")
    
    return generate_word_bytes(doc)
